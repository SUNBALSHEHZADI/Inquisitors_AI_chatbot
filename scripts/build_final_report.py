from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as PdfImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "final_project_report.md"
ARCHITECTURE = ROOT / "screenshots" / "system-architecture.png"
DOCX_OUTPUT = ROOT / "docs" / "final-project-report.docx"
PDF_OUTPUT = ROOT / "docs" / "final-project-report.pdf"


def parse_markdown():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    blocks = []
    index = 0

    while index < len(lines):
        line = lines[index]
        if not line.strip():
            index += 1
            continue
        if line.startswith("```"):
            code = []
            index += 1
            while index < len(lines) and not lines[index].startswith("```"):
                code.append(lines[index])
                index += 1
            blocks.append(("code", "\n".join(code)))
            index += 1
            continue
        if line.startswith("|"):
            rows = []
            while index < len(lines) and lines[index].startswith("|"):
                row = [cell.strip() for cell in lines[index].strip("|").split("|")]
                if not all(set(cell) <= set("-: ") for cell in row):
                    rows.append(row)
                index += 1
            blocks.append(("table", rows))
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            blocks.append((f"h{level}", line[level:].strip()))
            index += 1
            continue
        if line.startswith("- "):
            items = []
            while index < len(lines) and lines[index].startswith("- "):
                items.append(lines[index][2:].strip())
                index += 1
            blocks.append(("bullets", items))
            continue
        if re.match(r"^\d+\. ", line):
            items = []
            while index < len(lines) and re.match(r"^\d+\. ", lines[index]):
                items.append(re.sub(r"^\d+\. ", "", lines[index]).strip())
                index += 1
            blocks.append(("numbers", items))
            continue
        paragraph = [line]
        index += 1
        while index < len(lines) and lines[index].strip() and not lines[index].startswith(("#", "- ", "|", "```")) and not re.match(r"^\d+\. ", lines[index]):
            paragraph.append(lines[index])
            index += 1
        blocks.append(("p", " ".join(paragraph)))
    return blocks


def clean_inline(text):
    text = re.sub(r"!\[([^]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "").replace("`", "")
    return text


def build_docx(blocks):
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.color.rgb = RGBColor(16, 43, 58)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.color.rgb = RGBColor(7, 155, 180)

    for kind, value in blocks:
        if kind == "h1":
            paragraph = document.add_heading(clean_inline(value), level=0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "h2":
            document.add_heading(clean_inline(value), level=1)
        elif kind == "h3":
            document.add_heading(clean_inline(value), level=2)
        elif kind == "p":
            paragraph = document.add_paragraph(clean_inline(value))
            if value.startswith("**Project:") or value.startswith("**Team") or value.startswith("**Teammate") or value.startswith("**GitHub") or value.startswith("**Date"):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "bullets":
            for item in value:
                document.add_paragraph(clean_inline(item), style="List Bullet")
        elif kind == "numbers":
            for item in value:
                document.add_paragraph(clean_inline(item), style="List Number")
        elif kind == "code":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(value)
            run.font.name = "Consolas"
            run.font.size = Pt(8)
        elif kind == "table":
            if not value:
                continue
            table = document.add_table(rows=1, cols=len(value[0]))
            table.style = "Light Shading Accent 1"
            for cell, text in zip(table.rows[0].cells, value[0]):
                cell.text = clean_inline(text)
            for row in value[1:]:
                cells = table.add_row().cells
                for cell, text in zip(cells, row):
                    cell.text = clean_inline(text)
        if kind == "p" and "../screenshots/system-architecture.png" in value:
            document.add_picture(str(ARCHITECTURE), width=Inches(6.7))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save(DOCX_OUTPUT)


def pdf_text(text):
    text = clean_inline(text)
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return text


def build_pdf(blocks):
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="ReportTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=24, leading=29, alignment=TA_CENTER, textColor=colors.HexColor("#102b3a"), spaceAfter=14))
    styles.add(ParagraphStyle(name="ReportSubtitle", parent=styles["Normal"], fontSize=11, leading=16, alignment=TA_CENTER, textColor=colors.HexColor("#65777b"), spaceAfter=8))
    styles.add(ParagraphStyle(name="ReportH1", parent=styles["Heading1"], fontSize=16, leading=20, textColor=colors.HexColor("#102b3a"), spaceBefore=18, spaceAfter=8))
    styles.add(ParagraphStyle(name="ReportH2", parent=styles["Heading2"], fontSize=12, leading=16, textColor=colors.HexColor("#079bb4"), spaceBefore=12, spaceAfter=5))
    styles.add(ParagraphStyle(name="ReportBody", parent=styles["BodyText"], fontSize=9.5, leading=14, spaceAfter=7))
    styles.add(ParagraphStyle(name="ReportBullet", parent=styles["BodyText"], fontSize=9.5, leading=14, leftIndent=15, firstLineIndent=-8, bulletIndent=5, spaceAfter=3))
    styles.add(ParagraphStyle(name="ReportCode", parent=styles["Code"], fontName="Courier", fontSize=7.5, leading=10, backColor=colors.HexColor("#f1f6f4"), borderPadding=6, spaceAfter=8))

    story = []
    for kind, value in blocks:
        if kind == "h1":
            story.append(Paragraph(pdf_text(value), styles["ReportTitle"]))
        elif kind == "h2":
            story.append(Paragraph(pdf_text(value), styles["ReportH1"]))
        elif kind == "h3":
            story.append(Paragraph(pdf_text(value), styles["ReportH2"]))
        elif kind == "p":
            story.append(Paragraph(pdf_text(value), styles["ReportBody"]))
            if "../screenshots/system-architecture.png" in value:
                story.append(PdfImage(str(ARCHITECTURE), width=6.8 * inch, height=4.165 * inch))
                story.append(Spacer(1, 8))
        elif kind == "bullets":
            for item in value:
                story.append(Paragraph(pdf_text(item), styles["ReportBullet"], bulletText="•"))
        elif kind == "numbers":
            for number, item in enumerate(value, start=1):
                story.append(Paragraph(pdf_text(item), styles["ReportBullet"], bulletText=f"{number}."))
        elif kind == "code":
            story.append(Paragraph(pdf_text(value).replace("\n", "<br/>"), styles["ReportCode"]))
        elif kind == "table":
            if not value:
                continue
            table_data = [[Paragraph(pdf_text(cell), styles["ReportBody"]) for cell in row] for row in value]
            table = Table(table_data, repeatRows=1, hAlign="LEFT")
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#14516a")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d7e5e2")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.extend([Spacer(1, 5), table, Spacer(1, 8)])

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#65777b"))
        canvas.drawString(0.7 * inch, 0.4 * inch, "Inquisitors AI Assistant - Final Project Report")
        canvas.drawRightString(7.7 * inch, 0.4 * inch, f"Page {doc.page}")
        canvas.restoreState()

    document = SimpleDocTemplate(str(PDF_OUTPUT), pagesize=A4, rightMargin=0.65 * inch, leftMargin=0.65 * inch, topMargin=0.65 * inch, bottomMargin=0.65 * inch)
    document.build(story, onFirstPage=footer, onLaterPages=footer)


if __name__ == "__main__":
    report_blocks = parse_markdown()
    build_docx(report_blocks)
    build_pdf(report_blocks)
    print(f"Created {DOCX_OUTPUT}")
    print(f"Created {PDF_OUTPUT}")
